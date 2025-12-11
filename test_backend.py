import os
import sys

# Add backend to path
sys.path.append("/mnt/area51/Projects/ai-resume-parser/smart-resume-analyzer/backend")

from parser import parse_resume
from scorer import calculate_ats_score
from database import db

def test_backend():
    # 1. Create a dummy resume file for testing
    dummy_resume_path = "test_resume.docx"
    try:
        import docx
        doc = docx.Document()
        doc.add_paragraph("John Doe")
        doc.add_paragraph("johndoe@example.com | 555-0102")
        doc.add_heading("Summary", level=1)
        doc.add_paragraph("Experienced Software Engineer with 5 years in Python and Java. Increased efficiency by 20%.")
        doc.add_heading("Skills", level=1)
        doc.add_paragraph("Python, Java, SQL, Machine Learning, Communication")
        doc.add_heading("Experience", level=1)
        doc.add_paragraph("Software Engineer at Tech Corp. 2020-Present.")
        doc.add_heading("Education", level=1)
        doc.add_paragraph("BS Computer Science, University of Tech.")
        doc.save(dummy_resume_path)
        print(f"Created dummy resume at {dummy_resume_path}")
    except ImportError:
        print("python-docx not installed directly in this env? Skipping creation if so.")
        return

    # 2. Test Parser
    print("\n--- Testing Parser ---")
    data = parse_resume(dummy_resume_path)
    print("Parsed Data:", data)
    
    if "error" in data:
        print("Parser failed!")
        return

    # 3. Test Scorer
    print("\n--- Testing Scorer ---")
    score_data = calculate_ats_score(data)
    print("Score Data:", score_data)

    # 4. Test Database
    print("\n--- Testing Database ---")
    try:
        resume_id = db.insert_resume({**data, **score_data})
        print(f"Inserted resume with ID: {resume_id}")
        
        fetched = db.get_resume(resume_id)
        if fetched:
            print("Successfully fetched resume from DB.")
        else:
            print("Failed to fetch resume.")
    except Exception as e:
        print(f"Database error (expected if no local mongo): {e}")

    # Cleanup
    if os.path.exists(dummy_resume_path):
        os.remove(dummy_resume_path)

if __name__ == "__main__":
    test_backend()
